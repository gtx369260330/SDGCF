from pathlib import Path

from docx import Document
from docx.shared import RGBColor


RED = RGBColor(192, 0, 0)


def newest_manuscript(docs_dir: Path) -> Path:
    candidates = sorted(
        docs_dir.glob("*.docx"),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    for path in candidates:
        if "论文修订稿" in path.name and "红色修订版" not in path.name:
            return path
    raise FileNotFoundError("No manuscript draft matching docs/*论文修订稿*.docx was found.")


def red_run(paragraph, text: str):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.color.rgb = RED
    return run


def replace_first(doc: Document, needle: str, replacement: str, log: list[str]) -> bool:
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        if needle in paragraph.text:
            red_run(paragraph, replacement)
            log.append(f"P{i:03d}: {needle[:70]}")
            return True
    log.append(f"NOT FOUND: {needle[:70]}")
    return False


def replace_exact_heading(doc: Document, old: str, new: str, log: list[str]) -> bool:
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        if paragraph.text.strip() == old:
            red_run(paragraph, new)
            log.append(f"P{i:03d}: {old}")
            return True
    log.append(f"NOT FOUND HEADING: {old}")
    return False


def main():
    root = Path.cwd()
    docs_dir = root / "docs"
    src = newest_manuscript(docs_dir)
    out = docs_dir / "论文修订稿_红色修订版.docx"
    report = root / "results" / "reports" / "redline_revision_log.txt"
    report.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(src)
    log = [f"Source: {src}", f"Output: {out}", ""]

    replacements = [
        (
            "Results: On the held-out test set, SDGCF achieved an accuracy of 0.7824",
            "Results: On the held-out test set, SDGCF achieved an accuracy of 0.7824, balanced accuracy of 0.7498, Macro-F1 of 0.7444, weighted F1 of 0.7947, Cohen's kappa of 0.7095, Matthews correlation coefficient (MCC) of 0.7129, macro-AUROC of 0.9494 and macro-AUPRC of 0.8030. Compared with the graph-free ablation (Macro-F1 = 0.7324), SDGCF showed a 0.0120 absolute Macro-F1 gain, indicating a modest benefit associated with sample-wise dynamic graph interaction. In robustness experiments, SDGCF achieved the highest mean perturbed Macro-F1 among the evaluated multimodal models (0.6632) and led under missing Fpz-Cz, missing Pz-Oz and random single-channel missingness. It was not the top model under every corruption condition, including missing EOG, two-channel missingness and several noisy-channel settings. These findings support condition-dependent robustness rather than uniform superiority across all perturbations. The learned mean graph attention matrix was asymmetric across modalities, suggesting non-uniform EEG-EOG aggregation rather than simple average fusion.",
        ),
        (
            "Conclusion: SDGCF provides a compact and inspectable framework for EEG-EOG sleep staging",
            "Conclusion: SDGCF provides a compact and inspectable framework for EEG-EOG sleep staging by representing physiological channels as modality nodes and learning sample-wise dynamic graph interactions. The model achieved the strongest held-out test Macro-F1 among the evaluated models and showed competitive, condition-dependent robustness. Multi-seed experiments, fixed-graph comparisons, external validation and broader ablation studies remain necessary before making strong claims about generalization or clinical deployment.",
        ),
        (
            "The main contributions of this study are as follows:",
            "The main contributions of this study are summarized as follows:",
        ),
        (
            "We implement an EEG-EOG sleep-staging architecture that represents three physiological channels as modality nodes rather than treating them as anonymous stacked channels.",
            "We propose an EEG-EOG sleep-staging architecture that represents three physiological channels as modality nodes rather than treating them as anonymous stacked channels.",
        ),
        (
            "We design a multi-scale temporal encoding and dynamic graph attention pipeline for sample-wise cross-modality interaction among EEG Fpz-Cz, EEG Pz-Oz and horizontal EOG.",
            "We combine independent multi-scale temporal encoders with sample-wise dynamic graph attention to model cross-modality interactions among EEG Fpz-Cz, EEG Pz-Oz and horizontal EOG.",
        ),
        (
            "We train the model with a hybrid objective combining class-imbalance-aware classification terms, auxiliary modality supervision and graph regularization.",
            "We train the model with a hybrid objective that combines class-imbalance-aware classification losses, auxiliary modality supervision and graph regularization.",
        ),
        (
            "The project stores preprocessed 30 s epochs. Inputs are normalized into [N, C, T] format when required.",
            "The project stores preprocessed 30 s epochs in [N, C, T] format. The final manuscript should specify the exact normalization procedure, including whether normalization was performed per epoch, per channel, per subject or using training-set statistics (需作者补充依据).",
        ),
        (
            "During training, lightweight channel augmentation was applied: additive Gaussian noise was added to input channels",
            "During training, lightweight channel augmentation was applied: additive Gaussian noise was added to input channels, and modality dropout randomly masked channels while preserving at least one channel per sample. These augmentations were used only during training. The exact augmentation probabilities and noise scale must be reported with the released training configuration (需作者补充依据). Robustness testing used deterministic corruption conditions, including missing Fpz-Cz, missing Pz-Oz, missing EOG, missing one random channel, missing two channels and additive noise to each individual channel.",
        ),
        (
            "Figure X. Experimental workflow and model structure of SDGCF.",
            "Figure 2. Experimental workflow and model structure of SDGCF.",
        ),
        (
            "Figure 2 summarizes the main comparison on the held-out test set.",
            "Figure 3 summarizes the main comparison on the held-out test set.",
        ),
        (
            "Figure 2. Comparative performance heatmap for SDGCF and baseline models.",
            "Figure 3. Comparative performance heatmap for SDGCF and baseline models.",
        ),
        (
            "Additional visual summaries of model ranking, metric profiles and the top-model radar plot are provided in Supplementary Figures S1 and Figure S2.",
            "Additional visual summaries of model ranking, metric profiles and the top-model radar plot are provided in Supplementary Figures S1 and S2.",
        ),
        (
            "Training loss, Validation loss, Validation accuracy and Validation macro-F1",
            "Figure 4 shows the training dynamics of SDGCF, including training loss, validation loss, validation accuracy and validation Macro-F1. The progressive reduction in training loss together with stable validation metrics indicates stable optimization without marked late-stage degradation. Training dynamics for the ablation models and comparator algorithms are provided in the supplementary figures.",
        ),
        (
            "Figure 3. Training dynamics of SDGCF.",
            "Figure 4. Training dynamics of SDGCF.",
        ),
        (
            "Figure 4 presents the per-class performance of SDGCF.",
            "Figure 5 presents the per-class performance of SDGCF.",
        ),
        (
            "Figure 4. Per-class classification performance of SDGCF on the held-out test set.",
            "Figure 5. Per-class classification performance of SDGCF on the held-out test set.",
        ),
        (
            "As shown in Figure 5, the area under the curve (AUC) demonstrates strong discriminative ability",
            "As shown in Figure 6, the area under the curve (AUC) demonstrates strong discriminative ability for W (0.988), N3 (0.978), REM (0.967), and N2 (0.944), with comparatively lower separability for N1 (0.871). The diagonal dashed line represents random classification. Overall, these curves indicate stronger one-vs-rest separability for W, N2, N3 and REM than for N1.",
        ),
        (
            "Figure 5. One-vs-rest receiver operating characteristic curves for SDGCF on the held-out test set.",
            "Figure 6. One-vs-rest receiver operating characteristic curves for SDGCF on the held-out test set.",
        ),
        (
            "In Figure 6, average precision (AP) is highest for W",
            "In Figure 7, average precision (AP) is highest for W (0.979) and remains high for N2 (0.877), REM (0.861), and N3 (0.846), whereas N1 shows reduced AP (0.452), reflecting the difficulty of detecting this transitional and under-represented sleep stage.",
        ),
        (
            "Figure 6. One-vs-rest precision-recall curves for SDGCF on the held-out test set.",
            "Figure 7. One-vs-rest precision-recall curves for SDGCF on the held-out test set.",
        ),
        (
            "Figure 7 shows the raw confusion matrix of SDGCF on the test set.",
            "Figure 8 shows the normalized confusion matrix of SDGCF on the test set. Major error directions included W to N1, N1 to N2, N2 to N1, N3 to N2 and REM to N1. These patterns indicate that most errors occurred between neighboring or physiologically related stages.",
        ),
        (
            "Figure 7. Normalized confusion matrix of SDGCF on the held-out test set.",
            "Figure 8. Normalized confusion matrix of SDGCF on the held-out test set.",
        ),
        (
            "Values indicate the proportion of epochs from each true sleep stage assigned to each predicted stage.",
            "Values indicate the proportion of epochs from each true sleep stage assigned to each predicted stage. Misclassifications were concentrated between physiologically adjacent or transitional stages, particularly N1 with N2/REM and N3 with N2.",
        ),
        (
            "Representative hypnogram comparison",
            "To further examine temporal agreement between predicted and reference sleep-stage sequences, we performed a representative hypnogram comparison, as shown in Figure 9. Ground-truth and predicted sleep-stage sequences are plotted across consecutive 30 s epochs. The overall alignment across sustained sleep-stage periods indicates that SDGCF captures large-scale sleep architecture, whereas short discrepancies are mainly observed around stage transitions.",
        ),
        (
            "Figure 8. Representative hypnogram comparison for subject SC473.",
            "Figure 9. Representative hypnogram comparison for subject SC473.",
        ),
        (
            "The mean graph attention matrix on the SDGCF test set is shown in Figure 9.",
            "The mean graph attention matrix on the SDGCF test set is shown in Figure 10.",
        ),
        (
            "Figure 9. Mean dynamic graph attention matrix learned by SDGCF on the test set.",
            "Figure 10. Mean dynamic graph attention matrix learned by SDGCF on the test set.",
        ),
        (
            "Figure 10 shows that, attention matrices are grouped by ground-truth sleep stage",
            "Figure 11 shows stage-wise attention matrices grouped by ground-truth sleep stage, illustrating that cross-modality dependencies vary across W, N1, N2, N3 and REM epochs.",
        ),
        (
            "Figure 10. Sleep-stage-specific dynamic graph attention patterns learned by SDGCF.",
            "Figure 11. Sleep-stage-specific dynamic graph attention patterns learned by SDGCF.",
        ),
        (
            "Figure 11. Principal component analysis visualization of fused SDGCF feature embeddings.",
            "Figure 12. Principal component analysis visualization of fused SDGCF feature embeddings.",
        ),
        (
            "As shown Figure 12, in the current all experiments robustness summary",
            "As shown in Figure 13, in the finalized robustness experiment summary used for this manuscript (需作者确认最终结果源), SDGCF achieved Macro-F1 0.7444 under clean input, ranking first and exceeding the graph-free ablation (0.7324), Concat Transformer (0.7238) and Simple-Concatenation CNN (0.7028). Under perturbed inputs, SDGCF was the best model when Fpz-Cz was missing (0.6914), when Pz-Oz was missing (0.7222) and under random single-modality missingness (0.6985).",
        ),
        (
            "The other analyze are shown in Figure S11-S17.",
            "Additional robustness analyses are provided in Supplementary Figures S11-S17.",
        ),
        (
            "Figure 12. Robustness of SDGCF and multimodal baselines under channel missingness and additive noise.",
            "Figure 13. Robustness of SDGCF and multimodal baselines under channel missingness and additive noise.",
        ),
        (
            "A three-channel EEG-EOG model is relevant for simplified PSG analysis and may be useful",
            "A three-channel EEG-EOG model is relevant for simplified PSG analysis and may be useful in reduced-montage sleep-monitoring research settings.",
        ),
        (
            "N1 was the weakest stage, with F1-score 0.4911.",
            "N1 was the weakest stage, with an F1-score of 0.4911. This result is consistent with known challenges in sleep staging: N1 is relatively rare, often transitional and can share features with W, N2 and REM. In the confusion matrix, N1 was frequently confused with N2, and REM was frequently confused with N1. Future work should specifically address N1 by using sequence context, contrastive objectives, targeted data augmentation, calibrated loss functions or expert-rule-guided feature constraints.",
        ),
        (
            "Ethics approval and consent to participate: This study used a publicly available de-identified Sleep-EDF-derived dataset.",
            "Ethics approval and consent to participate: This study used a publicly available de-identified Sleep-EDF-derived dataset. Additional institutional approval and consent requirements should be confirmed according to local institutional and target-journal policies before submission (需作者确认).",
        ),
        (
            "Availability of data and materials: The original Sleep-EDF Expanded database is publicly available from PhysioNet.",
            "Availability of data and materials: The original Sleep-EDF Expanded database is publicly available from PhysioNet. The processed arrays, code and experiment scripts should be made available at a repository URL or provided upon reasonable request, depending on the target journal policy (需作者补充 repository URL and access conditions).",
        ),
        (
            "Competing interests: The authors declare that they have no competing interests",
            "Competing interests: The authors declare that they have no competing interests (需作者确认).",
        ),
        (
            "Funding: No specific funding was reported for this draft",
            "Funding: No specific funding was reported for this draft (需作者确认并按实际项目补充).",
        ),
        (
            "Authors' contributions: Author contributions should be completed",
            "Authors' contributions: Author contributions should be completed using the target journal's CRediT taxonomy, including conceptualization, methodology, software, validation, formal analysis, writing-original draft, writing-review and editing, supervision and funding acquisition as applicable (需作者补充依据).",
        ),
        (
            "American Academy of Sleep Medicine. The AASM Manual for the Scoring of Sleep and Associated Events",
            "American Academy of Sleep Medicine. The AASM Manual for the Scoring of Sleep and Associated Events: Rules, Terminology and Technical Specifications. Darien, IL: American Academy of Sleep Medicine. Version and edition year should be completed according to the target journal requirement (需作者补充依据).",
        ),
    ]

    for needle, replacement in replacements:
        replace_first(doc, needle, replacement, log)

    replace_exact_heading(doc, "5.33 Interpretation of attention", "5.3 Interpretation of attention", log)

    doc.save(out)
    report.write_text("\n".join(log), encoding="utf-8")
    print(out)
    print(report)


if __name__ == "__main__":
    main()
