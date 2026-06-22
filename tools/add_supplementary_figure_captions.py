from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


CAPTIONS = [
    (
        "Model comparison by Macro-F1 across evaluated models.",
        "The bar plot ranks SDGCF and comparator methods using the reported Macro-F1, providing a supplementary overview of class-balanced model performance.",
    ),
    (
        "Multi-metric radar comparison of SDGCF and leading comparator models.",
        "The radar plot summarizes complementary metrics, including accuracy, weighted F1, macro precision, macro recall and Macro-F1, to compare overall and class-balanced performance profiles.",
    ),
    (
        "Ablation impact of dynamic graph attention.",
        "The plot reports the Macro-F1 difference between full SDGCF and the graph-free SDGCF variant, illustrating the additional contribution associated with sample-wise dynamic graph interaction.",
    ),
    (
        "Training dynamics of SDGCF without dynamic graph attention.",
        "The curves provide the optimization reference for the graph-free ablation and support comparison with the full SDGCF architecture.",
    ),
    (
        "Performance summary of the Random Forest baseline.",
        "Because this classical baseline is not optimized epoch by epoch in the same way as neural models, the figure reports its final training and validation metrics as a compact summary.",
    ),
    (
        "Performance summary of the XGBoost baseline.",
        "The single-summary view reports the final training and validation metrics used to compare this boosted-tree classifier with neural and multimodal models.",
    ),
    (
        "Training dynamics of the Modality Transformer baseline.",
        "The curves show the optimization trajectory and validation performance of the Transformer-based fusion comparator under the same evaluation framework.",
    ),
    (
        "Performance summary of the Logistic Regression baseline.",
        "The figure reports the final training and validation metrics for this classical linear classifier, providing a low-complexity reference point.",
    ),
    (
        "Performance summary of the Linear SVM baseline.",
        "The figure reports the final training and validation metrics for the linear support vector machine comparator.",
    ),
    (
        "Training dynamics of the Simple-Concatenation CNN baseline.",
        "The curves show the behavior of a model that fuses channels by direct concatenation without explicit sample-wise graph interaction.",
    ),
    (
        "Training dynamics of the single-channel EEG Fpz-Cz baseline.",
        "The curves summarize model optimization when only the frontal EEG derivation is used as input.",
    ),
    (
        "Training dynamics of the single-channel EOG baseline.",
        "The curves summarize model optimization when only horizontal EOG is used, highlighting the standalone contribution and limitation of ocular information.",
    ),
    (
        "Training dynamics of the single-channel EEG Pz-Oz baseline.",
        "The curves summarize model optimization when only the parietal-occipital EEG derivation is used as input.",
    ),
    (
        "Robustness comparison under missing and noisy modality conditions.",
        "The line plot compares Macro-F1 across clean input, missing-channel settings and additive-noise settings for SDGCF and multimodal baselines.",
    ),
    (
        "Relative performance drop under modality corruption.",
        "The plot shows the percentage Macro-F1 decrease relative to each model's clean-input performance, summarizing corruption-specific sensitivity.",
    ),
    (
        "Heatmap of relative Macro-F1 drop under modality corruption.",
        "Darker cells indicate larger performance degradation relative to the corresponding clean-input condition, allowing direct comparison of corruption sensitivity across models.",
    ),
    (
        "Heatmap of absolute Macro-F1 under missing and noisy modality conditions.",
        "The heatmap compares retained classification performance across corruption scenarios and complements the relative-drop analysis.",
    ),
]


def iter_image_blobs_in_document_order(doc: Document):
    """Yield image relationship targets in paragraph order."""
    seen = []
    for para in doc.paragraphs:
        for blip in para._p.xpath(".//a:blip"):
            rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if not rid:
                continue
            rel = doc.part.rels[rid]
            seen.append((rid, rel.target_ref, rel.target_part.blob))
    return seen


def set_paragraph_font(paragraph, size_pt: float, bold: bool = False, color: RGBColor | None = None):
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color


def add_caption(paragraph, fig_no: int, title: str, explanation: str):
    label = paragraph.add_run(f"Figure S{fig_no}. ")
    label.bold = True
    title_run = paragraph.add_run(f"{title} ")
    title_run.bold = True
    paragraph.add_run(explanation)
    set_paragraph_font(paragraph, 9.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.05


def set_picture_alt_text(paragraph, fig_no: int, title: str):
    for doc_pr in paragraph._p.xpath(".//wp:docPr"):
        doc_pr.set("title", f"Figure S{fig_no}")
        doc_pr.set("descr", title)


def build_docx(input_docx: Path, output_docx: Path, image_dir: Path):
    source = Document(str(input_docx))
    images = iter_image_blobs_in_document_order(source)
    if len(images) != len(CAPTIONS):
        raise RuntimeError(f"Expected {len(CAPTIONS)} images, found {len(images)} in document order.")

    if image_dir.exists():
        shutil.rmtree(image_dir)
    image_dir.mkdir(parents=True, exist_ok=True)

    image_paths = []
    for idx, (_, target, blob) in enumerate(images, start=1):
        suffix = Path(target).suffix or ".png"
        image_path = image_dir / f"figure_s{idx:02d}{suffix}"
        image_path.write_bytes(blob)
        image_paths.append(image_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Supplementary Figures")
    title_run.font.name = "Times New Roman"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.paragraph_format.space_after = Pt(12)

    usable_width_in = (
        section.page_width - section.left_margin - section.right_margin
    ) / 914400
    figure_width = min(6.6, usable_width_in)

    for idx, (image_path, (caption_title, explanation)) in enumerate(zip(image_paths, CAPTIONS), start=1):
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.paragraph_format.space_before = Pt(0)
        pic_para.paragraph_format.space_after = Pt(2)
        pic_para.add_run().add_picture(str(image_path), width=Inches(figure_width))
        set_picture_alt_text(pic_para, idx, caption_title)

        cap_para = doc.add_paragraph()
        add_caption(cap_para, idx, caption_title, explanation)

        if idx != len(CAPTIONS):
            doc.add_page_break()

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return len(images)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--input", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    parser.add_argument("--image-dir", required=True, type=Path)
    args = parser.parse_args()

    count = build_docx(args.input, args.output, args.image_dir)
    print(f"Wrote {args.output} with {count} figures and {len(CAPTIONS)} captions.")


if __name__ == "__main__":
    main()
